#!/usr/bin/env python3
"""
REGOLO — COMPILATORE DI REGOLAMENTI (L1) — v0.2, AI reale + schema Pydantic.

Novità v0.2 (decisioni dagli spike del 2026-06-11):
  * struttura imposta dallo schema Pydantic (`schema_contratto.py`), non descritta
    nel prompt — validatori di business inclusi (coerenza checklist↔escalation)
  * chiamata via INSTRUCTOR (`instructor.from_genai`) con retry automatici sugli
    errori di validazione — scelto sul `response_schema` nativo perché preserva
    le escalation di ambiguità (vedi spikes/instructor/RISULTATO.md)
  * input anche PDF e DOCX: i regolamenti reali si caricano direttamente

Output in compilato_ai/:
  gara_compilata_ai.yaml        contratto proposto (in attesa di firma umana)
  compilazione_report_ai.json   report clausole + checklist + metadati run
  confronto_v1.json             auto-verifica vs contratto v1 approvato (se presente)

Uso:
  venv/bin/python mockup/01_contratto/compilatore.py
  venv/bin/python mockup/01_contratto/compilatore.py --regolamento percorso.pdf
  venv/bin/python mockup/01_contratto/compilatore.py --provider ollama
"""

import argparse
import json
import sys
import time
from pathlib import Path

import yaml

sys.path.insert(0, str(Path(__file__).parent))
from schema_contratto import SOGLIA_CONFIDENZA, Compilazione  # noqa: E402

QUI = Path(__file__).parent
ROOT = QUI.parent.parent
REGOLAMENTO_DEFAULT = QUI.parent / "00_regolamento" / "REGOLAMENTO_SPAZIO_ALLA_META_2026.md"
OUT = QUI / "compilato_ai"


def carica_env() -> dict:
    """Parser minimale del .env di progetto (niente dipendenze extra)."""
    env = {}
    f = ROOT / ".env"
    if f.exists():
        for riga in f.read_text().splitlines():
            riga = riga.strip()
            if riga and not riga.startswith("#") and "=" in riga:
                k, v = riga.split("=", 1)
                env[k.strip()] = v.strip().strip('"').strip("'")
    return env


# ---------------------------------------------------------------------------
# Estrazione testo: i regolamenti veri arrivano in PDF/DOCX
# ---------------------------------------------------------------------------

def estrai_testo(percorso: Path) -> str:
    suffisso = percorso.suffix.lower()
    if suffisso in (".md", ".txt"):
        return percorso.read_text()
    if suffisso == ".pdf":
        from pypdf import PdfReader
        pagine = [p.extract_text() or "" for p in PdfReader(str(percorso)).pages]
        return "\n\n".join(pagine)
    if suffisso == ".docx":
        import docx
        d = docx.Document(str(percorso))
        blocchi = [p.text for p in d.paragraphs]
        for tabella in d.tables:  # le tabelle (cluster, premi) sono spesso decisive
            for riga in tabella.rows:
                blocchi.append(" | ".join(c.text.strip() for c in riga.cells))
        return "\n".join(blocchi)
    raise SystemExit(f"formato non supportato: {suffisso} (attesi: .md .txt .pdf .docx)")


# ---------------------------------------------------------------------------
# Prompt: SOLO regole e checklist — la struttura la impone lo schema Pydantic
# ---------------------------------------------------------------------------

PROMPT = """Sei il Compilatore di Regolamenti di REGOLO (sistema loyalty di H&A Motivation).
Trasformi il regolamento di un'operazione a premi in un Contratto di Gara: una specifica
strutturata ed ESEGUIBILE da un motore di calcolo deterministico.

REGOLE TASSATIVE:
1. Estrai SOLO ciò che il regolamento dice. Non inventare valori, non "completare" regole.
2. Ogni elemento del contratto cita la clausola di origine in fonte_clausola (es. "art. 5.1").
3. Per ogni clausola interpretata, una voce di report con confidenza 0.0-1.0.
4. CHECKLIST AMBIGUITÀ — è il tuo compito più importante (errori qui muovono denaro).
   Percorri i 5 controlli uno per uno e registra l'esito:
   a) decorrenze_finestre_temporali: per ogni "entro N giorni/mesi", la decorrenza è
      specificata (firma? attivazione? caricamento)? Se no → ambiguita_trovata.
   b) ricorrenza_bonus: per ogni bonus/premio, è ESPLICITO se una tantum o ricorrente?
   c) cap_lordo_o_netto: i massimali si applicano al lordo o al netto degli storni?
   d) grandezze_definite: tutte le grandezze citate in condizioni/formule sono definite?
   e) termini_vaghi: termini vaghi con effetti sul calcolo?
   Ogni ambiguità trovata va anche: nel report come clausola stato "da_verificare" con
   motivo_escalation, e in contratto.clausole_da_verificare. Un'interpretazione
   "ragionevole ma non scritta" è un ERRORE, non un servizio.
5. Date ISO (YYYY-MM-DD), mesi "YYYY-MM", numeri come numeri.

REGOLAMENTO DA COMPILARE:
----------------------------------------
{REGOLAMENTO}
----------------------------------------"""


# ---------------------------------------------------------------------------
# Provider LLM
# ---------------------------------------------------------------------------

def chiama_gemini(prompt: str, env: dict) -> tuple[Compilazione, dict]:
    """Gemini via Vertex ADC, con instructor: schema + retry automatici."""
    import instructor
    from google import genai
    gclient = genai.Client(vertexai=True, project=env["GOOGLE_CLOUD_PROJECT"],
                           location=env.get("GOOGLE_CLOUD_LOCATION", "us-central1"))
    client = instructor.from_genai(gclient)
    modello = env.get("REGOLO_MODEL", "gemini-2.5-flash")
    t0 = time.time()
    risultato, raw = client.chat.completions.create_with_completion(
        model=modello,
        messages=[{"role": "user", "content": prompt}],
        response_model=Compilazione,
        max_retries=2)
    um = getattr(raw, "usage_metadata", None)
    meta = dict(provider="gemini-vertex (instructor)", modello=modello,
                durata_s=round(time.time() - t0, 1),
                token_input=getattr(um, "prompt_token_count", None),
                token_output=getattr(um, "candidates_token_count", None))
    return risultato, meta


def chiama_ollama(prompt: str, env: dict) -> tuple[Compilazione, dict]:
    """Fallback locale: schema JSON nel prompt + validazione Pydantic (2 tentativi)."""
    import urllib.request
    modello = env.get("REGOLO_MODEL_OLLAMA", "llama3.1")
    schema = json.dumps(Compilazione.model_json_schema(), ensure_ascii=False)
    contents = prompt + ("\n\nRispondi SOLO con un oggetto JSON valido conforme a questo "
                         f"JSON Schema:\n{schema}")
    t0, errore = time.time(), ""
    for _ in range(2):
        corpo = json.dumps(dict(model=modello, prompt=contents + errore, stream=False,
                                format="json", options=dict(temperature=0))).encode()
        req = urllib.request.Request("http://localhost:11434/api/generate", data=corpo,
                                     headers={"Content-Type": "application/json"})
        with urllib.request.urlopen(req, timeout=900) as r:
            grezzo = json.loads(r.read())["response"]
        try:
            ris = Compilazione.model_validate_json(grezzo)
            return ris, dict(provider="ollama", modello=modello,
                             durata_s=round(time.time() - t0, 1))
        except Exception as e:
            errore = f"\n\nIl tuo output precedente non valida: {str(e)[:300]}. Correggi."
    raise RuntimeError("Ollama: validazione fallita dopo 2 tentativi")


def ollama_disponibile() -> bool:
    import urllib.request
    try:
        with urllib.request.urlopen("http://localhost:11434/api/tags", timeout=3) as r:
            return len(json.loads(r.read()).get("models", [])) > 0
    except Exception:
        return False


# ---------------------------------------------------------------------------
# Auto-verifica: parametri AI vs contratto v1 approvato (ground truth)
# (funzioni riusate anche dagli spike — non cambiarne le firme)
# ---------------------------------------------------------------------------

SINONIMI_KPI = {
    "luce_gas": ["luce_gas", "commodity", "luce", "gas"],
    "fibra": ["fibra"],
    "fotovoltaico": ["fotovoltaico", "fv"],
    "wallbox": ["wallbox", "juicebox"],
    "clima": ["clima", "climatizzatore", "condizionatore"],
}


def _mecc_per_kpi(contratto: dict, kpi: str) -> dict | None:
    for nome in SINONIMI_KPI.get(kpi, [kpi]):
        for m in contratto.get("meccaniche", []):
            if m.get("kpi") == nome or nome in str(m.get("id", "")).lower():
                return m
    return None


def _bonus_con(contratto: dict, parole: list[str]) -> dict | None:
    for m in contratto.get("meccaniche", []):
        blob = (str(m.get("id", "")) + " " + str(m.get("condizione", ""))).lower()
        if "bonus" in (str(m.get("tipo", "")) + str(m.get("id", ""))).lower() \
           and all(p in blob for p in parole):
            return m
    return None


def _target_cluster(contratto: dict, cid: str):
    for d in contratto.get("cluster", {}).get("definizioni", []):
        if str(d.get("id", "")).upper() == cid:
            return d.get("target_mensile_commodity")
    return None


def estrai_parametri(c: dict) -> dict:
    """I parametri 'che muovono soldi': base del confronto AI vs approvato."""
    p = {}
    for kpi in ("luce_gas", "fibra", "fotovoltaico", "wallbox", "clima"):
        m = _mecc_per_kpi(c, kpi) or {}
        p[f"punti_{kpi}"] = m.get("punti")
    m_comm = _mecc_per_kpi(c, "luce_gas") or {}
    p["cap_commodity"] = m_comm.get("cap_mensile_punti")
    m_clima = _mecc_per_kpi(c, "clima") or {}
    molt = (m_clima.get("moltiplicatori") or [{}])[0]
    p["clima_moltiplicatore"] = molt.get("fattore")
    p["clima_mesi_moltiplicatore"] = sorted(str(x) for x in molt.get("periodi", [])) or None
    b_ob = _bonus_con(c, ["obiettivo"]) or _bonus_con(c, ["target"]) or {}
    p["bonus_obiettivo"] = b_ob.get("punti")
    b_tr = _bonus_con(c, ["fotovoltaico", "wallbox"]) or _bonus_con(c, ["tripletta"]) or {}
    p["bonus_tripletta"] = b_tr.get("punti")
    for cid in ("GOLD", "SILVER", "BRONZE"):
        p[f"target_{cid}"] = _target_cluster(c, cid)
    p["tasso_eur_per_punto"] = c.get("valorizzazione", {}).get("tasso_eur_per_punto")
    premi = c.get("classifica_finale", {}).get("premi_punti_extra", [])
    p["premi_classifica"] = sorted((x.get("punti") for x in premi), reverse=True) or None
    p["valido_dal"] = str(c.get("meta", {}).get("valido_dal", ""))[:10] or None
    p["valido_al"] = str(c.get("meta", {}).get("valido_al", ""))[:10] or None
    return p


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    ap = argparse.ArgumentParser(description="REGOLO — Compilatore L1 v0.2 (AI + schema Pydantic)")
    ap.add_argument("--regolamento", default=str(REGOLAMENTO_DEFAULT),
                    help="percorso .md/.txt/.pdf/.docx")
    ap.add_argument("--provider", choices=["auto", "gemini", "ollama"], default="auto")
    ap.add_argument("--out", default=str(OUT), help="cartella di output (default: compilato_ai del mockup)")
    args = ap.parse_args()
    out = Path(args.out)

    env = carica_env()
    percorso = Path(args.regolamento)
    testo = estrai_testo(percorso)
    prompt = PROMPT.replace("{REGOLAMENTO}", testo)

    provider = args.provider
    if provider == "auto":
        adc = Path.home() / ".config/gcloud/application_default_credentials.json"
        if env.get("GOOGLE_CLOUD_PROJECT") and adc.exists():
            provider = "gemini"
        elif ollama_disponibile():
            provider = "ollama"
        else:
            raise SystemExit("Nessun provider LLM disponibile: configurare Vertex ADC nel .env "
                             "o avviare Ollama con un modello installato.")

    print(f"▸ Compilo «{percorso.name}» ({len(testo)} caratteri) con provider: {provider}…")
    risultato, meta = (chiama_gemini if provider == "gemini" else chiama_ollama)(prompt, env)
    contratto = risultato.contratto.model_dump(exclude_none=True)
    clausole = [c.model_dump() for c in risultato.report]
    checklist = [v.model_dump() for v in risultato.checklist_ambiguita]

    out.mkdir(parents=True, exist_ok=True)
    with open(out / "gara_compilata_ai.yaml", "w") as f:
        f.write("# Contratto di Gara PROPOSTO DALL'AI — in attesa di review e approvazione umana.\n"
                f"# Generato da: {meta['provider']} / {meta['modello']} — NON usato dal motore\n"
                "# finché non viene approvato (HITL).\n")
        yaml.dump(contratto, f, allow_unicode=True, sort_keys=False)

    da_verificare = [c for c in clausole if c["stato"] == "da_verificare"]
    report_completo = dict(
        _descrizione="Report di compilazione PRODOTTO DALL'AI (L1 v0.2: Pydantic + instructor)",
        documento_fonte=percorso.name, esecuzione=meta,
        soglia_confidenza=SOGLIA_CONFIDENZA,
        statistiche=dict(clausole_analizzate=len(clausole),
                         escalation_verifica_umana=len(da_verificare)),
        checklist_ambiguita=checklist, clausole=clausole)
    (out / "compilazione_report_ai.json").write_text(
        json.dumps(report_completo, ensure_ascii=False, indent=1))

    # ----- auto-verifica vs contratto approvato (solo per la gara del mockup) ----
    v1_path = QUI / "gara_spazio_alla_meta.v1.yaml"
    if v1_path.exists() and percorso == REGOLAMENTO_DEFAULT:
        with open(v1_path) as f:
            v1 = yaml.safe_load(f)
        p_v1, p_ai = estrai_parametri(v1), estrai_parametri(contratto)
        confronto = [dict(parametro=k, approvato_v1=p_v1[k], proposto_ai=p_ai.get(k),
                          esito="OK" if p_ai.get(k) == p_v1[k] else "DIFF") for k in p_v1]
        conformi = sum(1 for r in confronto if r["esito"] == "OK")
        (out / "confronto_v1.json").write_text(json.dumps(dict(
            descrizione="Auto-verifica: parametri della compilazione AI vs contratto v1 approvato",
            conformi=conformi, totale=len(confronto), esecuzione=meta,
            righe=confronto), ensure_ascii=False, indent=1))
        print(f"\n  AUTO-VERIFICA — parametri vs contratto v1 approvato: {conformi}/{len(confronto)}")
        for r in confronto:
            segno = "✓" if r["esito"] == "OK" else "✗"
            extra = "" if r["esito"] == "OK" else f"   (AI: {r['proposto_ai']} | v1: {r['approvato_v1']})"
            print(f"   {segno} {r['parametro']}{extra}")

    print(f"\n✓ Compilazione in {meta['durata_s']}s "
          f"({meta.get('token_input', '?')} token in / {meta.get('token_output', '?')} out)")
    print(f"  Checklist ambiguità: {sum(1 for v in checklist if v['esito'] == 'ambiguita_trovata')} "
          f"ambiguità trovate su {len(checklist)} controlli")
    print(f"  ESCALATION A VERIFICA UMANA: {len(da_verificare)} clausole")
    for c in da_verificare:
        print(f"   ⚠ {c['articolo']} (conf. {c['confidenza']}): "
              f"{(c.get('motivo_escalation') or c['interpretazione'])[:110]}")
    print(f"\n  Output in {out}/ — il contratto AI passa ora dalla review umana (HITL).")


if __name__ == "__main__":
    main()
